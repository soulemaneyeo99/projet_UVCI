#!/usr/bin/env python3
"""Génère Script_BD_UVCI.docx — document Word académique propre."""

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

LOGO_MIN  = '/tmp/docx_imgs/14518ad006362d68c9948ef2b155814428506f51.png'
LOGO_UVCI = '/tmp/docx_imgs/1a1e88b7b37ea3fb5b85c07c44d804697606eecd.png'

# ─── couleurs ────────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x4E, 0x79)
BLUE   = RGBColor(0x2E, 0x75, 0xB6)
GREY10 = RGBColor(0xD9, 0xE2, 0xF3)   # header cellule
GREY05 = RGBColor(0xF2, 0xF5, 0xFA)   # ligne alternée
BLACK  = RGBColor(0x00, 0x00, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x37, 0x59, 0x23)   # commentaires SQL
KWORD  = RGBColor(0x00, 0x00, 0x80)   # mots-clés SQL bleu foncé
STRCLR = RGBColor(0x9B, 0x26, 0x00)   # chaînes SQL

# ─── helpers ─────────────────────────────────────────────────────────────────

def set_font(run, size=12, bold=False, italic=False, color=None, name='Times New Roman'):
    run.font.name        = name
    run.font.size        = Pt(size)
    run.font.bold        = bold
    run.font.italic      = italic
    if color:
        run.font.color.rgb = color

def para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=6,
             line_spacing=1.15, keep_with_next=False):
    pf = p.paragraph_format
    pf.alignment      = align
    pf.space_before   = Pt(space_before)
    pf.space_after    = Pt(space_after)
    pf.line_spacing   = Pt(line_spacing * 12)
    pf.keep_with_next = keep_with_next

def shade_cell(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}')
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   kwargs.get('val', 'single'))
        b.set(qn('w:sz'),    kwargs.get('sz',  '4'))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), kwargs.get('color', '000000'))
        tcBorders.append(b)
    tcPr.append(tcBorders)

def remove_borders(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_dot_toc_entry(doc, text, page_str, indent_cm=0, bold=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    if indent_cm:
        pf.left_indent = Cm(indent_cm)
    tab = OxmlElement('w:tabs')
    tb  = OxmlElement('w:tab')
    tb.set(qn('w:val'),    'right')
    tb.set(qn('w:leader'), 'dot')
    tb.set(qn('w:pos'),    '8640')
    tab.append(tb)
    p._p.pPr.append(tab)

    r1 = p.add_run(text)
    set_font(r1, 12, bold=bold)

    r2 = p.add_run('\t' + page_str)
    set_font(r2, 12, bold=bold)

def section_title(doc, num, title):
    """Titre de section principal (SECTION I. …)."""
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = True
    p.paragraph_format.space_after       = Pt(6)
    p.paragraph_format.space_before      = Pt(0)
    r = p.add_run(f'SECTION {num}. {title.upper()}')
    set_font(r, 14, bold=True, color=NAVY)

def subsection_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(title)
    set_font(r, 12, bold=True, color=BLUE)
    p.paragraph_format.keep_with_next = True

def hr(doc, color='2E75B6'):
    """Ligne horizontale via paragraphe avec bordure basse."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_text(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12,
             bold=False, color=None, sb=0, sa=6):
    p = doc.add_paragraph()
    para_fmt(p, align=align, space_before=sb, space_after=sa)
    r = p.add_run(text)
    set_font(r, size, bold=bold, color=color)
    return p

# ─── SQL colorizer ────────────────────────────────────────────────────────────

SQL_KEYWORDS = {
    'CREATE','TABLE','DROP','IF','EXISTS','CASCADE','BEGIN','COMMIT',
    'INSERT','INTO','VALUES','SET','INDEX','ON','UNIQUE','PRIMARY','KEY',
    'FOREIGN','REFERENCES','CONSTRAINT','CHECK','DEFAULT','NOT','NULL',
    'INTEGER','VARCHAR','SMALLINT','NUMERIC','BOOLEAN','DATE','TIMESTAMPTZ',
    'GENERATED','ALWAYS','AS','IDENTITY','UPDATE','DELETE','WHERE','TRUE','FALSE',
    'COMMENT','IS','IN','OR','AND','DISTINCT','SELECT','FROM',
}

def tokenize_sql(line):
    """Retourne liste de (texte, type) où type in comment|string|keyword|other."""
    # commentaire -- jusqu'à fin de ligne
    if re.match(r'\s*--', line):
        return [('comment', line)]
    tokens = []
    i = 0
    while i < len(line):
        # string littérale
        if line[i] == "'":
            j = i + 1
            while j < len(line):
                if line[j] == "'" and (j+1 < len(line) and line[j+1] == "'"):
                    j += 2
                elif line[j] == "'":
                    j += 1
                    break
                else:
                    j += 1
            tokens.append(('string', line[i:j]))
            i = j
        # mot (keyword ou identifiant)
        elif re.match(r'[A-Za-z_]', line[i]):
            j = i
            while j < len(line) and re.match(r'[A-Za-z0-9_]', line[j]):
                j += 1
            word = line[i:j]
            if word.upper() in SQL_KEYWORDS:
                tokens.append(('keyword', word))
            else:
                tokens.append(('other', word))
            i = j
        else:
            # autres caractères — agréger
            if tokens and tokens[-1][0] == 'other':
                tokens[-1] = ('other', tokens[-1][1] + line[i])
            else:
                tokens.append(('other', line[i]))
            i += 1
    return tokens

def add_sql_block(doc, sql_text):
    """Ajoute un bloc SQL colorisé (police Courier New 9pt) dans un cadre gris."""
    # bloc englobant : table 1×1 avec fond gris clair
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    shade_cell(cell, RGBColor(0xF5, 0xF5, 0xF5))
    set_cell_border(cell, val='single', sz='4', color='BFBFBF')

    # vider le paragraphe auto généré
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

    for raw_line in sql_text.strip('\n').split('\n'):
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.line_spacing = Pt(13)

        tokens = tokenize_sql(raw_line)
        for kind, text in tokens:
            r = p.add_run(text)
            r.font.name = 'Courier New'
            r.font.size = Pt(8.5)
            if kind == 'comment':
                r.font.color.rgb = GREEN
                r.font.italic    = True
            elif kind == 'keyword':
                r.font.color.rgb = KWORD
                r.font.bold      = True
            elif kind == 'string':
                r.font.color.rgb = STRCLR
            else:
                r.font.color.rgb = BLACK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── tableaux de description ─────────────────────────────────────────────────

def add_col_table(doc, table_name, rows_data, caption=None):
    """
    Tableau de description des colonnes d'une table SQL.
    rows_data : list of (colonne, type, contraintes, description)
    """
    tbl = doc.add_table(rows=1 + len(rows_data), cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # largeurs colonnes
    widths = [3.5, 3.0, 4.5, 5.5]
    for i, w in enumerate(widths):
        set_col_width(tbl, i, w)

    # en-tête
    headers = ['Colonne', 'Type SQL', 'Contraintes', 'Description']
    for i, h in enumerate(headers):
        cell = tbl.cell(0, i)
        shade_cell(cell, GREY10)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, 10, bold=True, color=NAVY)

    # données
    for ri, (col, typ, cst, desc) in enumerate(rows_data):
        bg = GREY05 if ri % 2 == 1 else WHITE
        vals = [col, typ, cst, desc]
        for ci, val in enumerate(vals):
            cell = tbl.cell(ri + 1, ci)
            if bg != WHITE:
                shade_cell(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            if ci == 0:
                set_font(r, 9.5, bold=True, name='Courier New', color=RGBColor(0x1F,0x4E,0x79))
            else:
                set_font(r, 9.5, name='Courier New' if ci == 1 else 'Times New Roman')

    if caption:
        cp = doc.add_paragraph(caption)
        cp.paragraph_format.space_before = Pt(2)
        cp.paragraph_format.space_after  = Pt(10)
        cp.runs[0].font.italic = True
        cp.runs[0].font.size   = Pt(10)
        cp.runs[0].font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ─── construction du document ─────────────────────────────────────────────────

doc = Document()

# ── marges A4 ────────────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_height      = Cm(29.7)
sec.page_width       = Cm(21.0)
sec.left_margin      = Cm(3.0)
sec.right_margin     = Cm(2.5)
sec.top_margin       = Cm(2.5)
sec.bottom_margin    = Cm(2.5)
sec.header_distance  = Cm(1.5)
sec.footer_distance  = Cm(1.5)

# ════════════════════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ════════════════════════════════════════════════════════════════════════════

logo_tbl = doc.add_table(rows=1, cols=2)
remove_borders(logo_tbl)
logo_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

lc = logo_tbl.cell(0, 0)
lc.width = Cm(7)
lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p_min = lc.paragraphs[0]
p_min.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_min = p_min.add_run()
run_min.add_picture(LOGO_MIN, width=Cm(5.5))

rc = logo_tbl.cell(0, 1)
rc.width = Cm(9)
rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p_uvci = rc.paragraphs[0]
p_uvci.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_uvci = p_uvci.add_run()
run_uvci.add_picture(LOGO_UVCI, width=Cm(5.5))

hr(doc, '1F4E79')

# Intitulé ministère
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after  = Pt(2)
r = p.add_run('MINISTÈRE DE L\'ENSEIGNEMENT SUPÉRIEUR ET DE LA RECHERCHE SCIENTIFIQUE')
set_font(r, 10, bold=True, color=NAVY)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run('Université Virtuelle de Côte d\'Ivoire — UVCI')
set_font(r2, 10, color=NAVY)

hr(doc, '2E75B6')

# Espace
for _ in range(4):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Type de document
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SCRIPT DE CRÉATION DE LA BASE DE DONNÉES')
set_font(r, 16, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(8)

# Titre principal
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('SYSTÈME DE GESTION DES HEURES PÉDAGOGIQUES')
set_font(r2, 20, bold=True, color=NAVY)
p2.paragraph_format.space_after = Pt(4)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('DES ENSEIGNANTS DE L\'UVCI')
set_font(r3, 20, bold=True, color=NAVY)
p3.paragraph_format.space_after = Pt(20)

# Ligne de séparation déco
hr(doc, '2E75B6')

# Bloc info technique
info_tbl = doc.add_table(rows=5, cols=2)
remove_borders(info_tbl)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

infos = [
    ('SGBD cible',   'PostgreSQL 14+'),
    ('Encodage',     'UTF-8'),
    ('Version',      '1.0'),
    ('Référence',    'PCT25-26_DAS-N°11'),
    ('Année',        '2025 – 2026'),
]
for i, (lbl, val) in enumerate(infos):
    c0 = info_tbl.cell(i, 0)
    c1 = info_tbl.cell(i, 1)
    c0.width = Cm(5)
    c1.width = Cm(7)
    r0 = c0.paragraphs[0].add_run(lbl + ' :')
    set_font(r0, 11, bold=True, color=NAVY)
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = c1.paragraphs[0].add_run(val)
    set_font(r1, 11, color=BLACK)
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

hr(doc, '2E75B6')

# Espace avant pied de page
for _ in range(6):
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(0)

# Auteurs
auteurs_data = [
    ['YEO Yanougui Souleymane',    'Développeur Full-Stack'],
    ['Collaborateur 2',            ''],
    ['Collaborateur 3',            ''],
]
encadreur = 'M. [Nom Prénom]'

auth_tbl = doc.add_table(rows=1, cols=2)
remove_borders(auth_tbl)
auth_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

ac0 = auth_tbl.cell(0, 0)
ac0.width = Cm(8)
p_a = ac0.paragraphs[0]
p_a.alignment = WD_ALIGN_PARAGRAPH.LEFT
ra0 = p_a.add_run('Réalisé par :')
set_font(ra0, 11, bold=True, color=NAVY)

ac1 = auth_tbl.cell(0, 1)
ac1.width = Cm(8)
p_b = ac1.paragraphs[0]
p_b.alignment = WD_ALIGN_PARAGRAPH.LEFT
rb0 = p_b.add_run('Encadreur :')
set_font(rb0, 11, bold=True, color=NAVY)

for nom, role in auteurs_data:
    rw = auth_tbl.add_row()
    c0 = rw.cells[0]
    c1 = rw.cells[1]
    c0.width = Cm(8)
    c1.width = Cm(8)
    r0 = c0.paragraphs[0].add_run('  ' + nom)
    set_font(r0, 11)
    c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = c1.paragraphs[0].add_run('  ' + encadreur if nom == auteurs_data[0][0] else '')
    set_font(r1, 11)
    c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    encadreur = ''

hr(doc, '1F4E79')

pf = doc.add_paragraph()
pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rf = pf.add_run('Groupe 11  ·  Année académique 2025–2026')
set_font(rf, 11, color=NAVY)

# Saut de page
doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# SOMMAIRE
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('SOMMAIRE')
set_font(r, 14, bold=True, color=NAVY)
hr(doc, '2E75B6')

toc_entries = [
    ('Introduction', '3', 0, True),
    ('Conventions de nommage', '3', 0, True),
    ('Section I.   Table USERS', '4', 0, True),
    ('Section II.  Table TEACHERS', '5', 0, True),
    ('Section III. Table COURSES', '6', 0, True),
    ('Section IV.  Table ACADEMIC_YEARS', '7', 0, True),
    ('Section V.   Table RESOURCES', '8', 0, True),
    ('Section VI.  Table ACTIVITIES', '9', 0, True),
    ('Section VII. Table COEFFICIENT_CONFIGS', '11', 0, True),
    ('Section VIII.Table QUOTAS_STATUTAIRES', '12', 0, True),
    ('Section IX.  Données de référence (INSERT)', '13', 0, True),
    ('Annexe A. Récapitulatif des tables', '14', 0, True),
    ('Annexe B. Matrice des contraintes', '15', 0, True),
    ('Annexe C. Schéma des relations', '15', 0, True),
]
for text, page, indent, bold in toc_entries:
    add_dot_toc_entry(doc, text, page, indent, bold)

doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# INTRODUCTION
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
r = p.add_run('INTRODUCTION')
set_font(r, 13, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
hr(doc)

add_text(doc, (
    'Le présent document constitue le script complet de définition du schéma relationnel '
    '(DDL — Data Definition Language) de la base de données du Système de Gestion des Heures '
    'Pédagogiques des Enseignants de l\'Université Virtuelle de Côte d\'Ivoire (UVCI). '
    'Ce script est destiné à être exécuté sur un serveur PostgreSQL 14 ou supérieur.'
))
add_text(doc, (
    'Il contient : (1) la suppression conditionnelle des tables existantes '
    '(DROP … IF EXISTS CASCADE), garantissant l\'idempotence du script ; '
    '(2) la création des huit tables du modèle relationnel avec leurs contraintes '
    'd\'intégrité, clés primaires, clés étrangères et index ; '
    '(3) l\'insertion des données de référence (barème officiel UVCI et quotas statutaires par défaut).'
))
add_text(doc, (
    'Le script est encadré dans une transaction explicite (BEGIN … COMMIT) : '
    'en cas d\'erreur, la totalité du script est annulée, préservant la cohérence de la base.'
))
add_text(doc, 'Commande d\'exécution :', sb=8, sa=2)

add_sql_block(doc, "    psql -U <utilisateur> -d <base_de_données> -f schema_uvci.sql")

# Conventions
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
r = p.add_run('CONVENTIONS DE NOMMAGE')
set_font(r, 13, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
hr(doc)

conv_rows = [
    ('Tables',          'Pluriel, snake_case',             'users, activities, academic_years'),
    ('Clés primaires',  'pk_<table>',                      'pk_users, pk_teachers'),
    ('Clés étrangères', 'fk_<table>_<colonne>',           'fk_teachers_user, fk_activities_resource'),
    ('Contraintes UNIQUE', 'uq_<table>_<colonnes>',       'uq_users_email, uq_coeff_niveau_type'),
    ('Contraintes CHECK', 'ck_<table>_<colonne>',         'ck_users_role, ck_teachers_statut'),
    ('Index',           'ix_<table>_<colonne>',            'ix_users_email, ix_activities_status'),
]

conv_tbl = doc.add_table(rows=1 + len(conv_rows), cols=3)
conv_tbl.style = 'Table Grid'
conv_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, h in enumerate(['Élément', 'Motif', 'Exemples']):
    cell = conv_tbl.cell(0, i)
    shade_cell(cell, GREY10)
    r = cell.paragraphs[0].add_run(h)
    set_font(r, 10, bold=True, color=NAVY)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

widths2 = [3.5, 4.0, 9.0]
for i, w in enumerate(widths2):
    set_col_width(conv_tbl, i, w)

for ri, (elem, motif, ex) in enumerate(conv_rows):
    bg = GREY05 if ri % 2 else WHITE
    for ci, val in enumerate([elem, motif, ex]):
        cell = conv_tbl.cell(ri + 1, ci)
        if bg != WHITE:
            shade_cell(cell, bg)
        r = cell.paragraphs[0].add_run(val)
        set_font(r, 9.5, name='Courier New' if ci > 0 else 'Times New Roman')

cp = doc.add_paragraph('Tableau 1. Conventions de nommage appliquées au schéma UVCI.')
cp.paragraph_format.space_before = Pt(2)
cp.paragraph_format.space_after  = Pt(10)
cp.runs[0].font.italic = True
cp.runs[0].font.size   = Pt(10)
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
# SECTIONS 1–8 : chaque table
# ════════════════════════════════════════════════════════════════════════════

sections = [

  # ── SECTION I : USERS ────────────────────────────────────────────────────
  dict(
    num='I', title='TABLE USERS — Comptes d\'authentification',
    description=(
        'La table USERS centralise les comptes d\'authentification de l\'application. '
        'Chaque utilisateur se voit attribuer un rôle RBAC (admin, secretary, teacher) '
        'qui détermine ses droits d\'accès à l\'API. '
        'Le mot de passe est systématiquement haché avec l\'algorithme bcrypt avant stockage : '
        'la valeur en clair n\'est jamais persistée.'
    ),
    sql='''\
-- ===================================================================
-- 1. UTILISATEURS (comptes d'authentification)
-- ===================================================================
CREATE TABLE users (
    id              INTEGER       GENERATED ALWAYS AS IDENTITY,
    email           VARCHAR(255)  NOT NULL,
    hashed_password VARCHAR(255)  NOT NULL,
    role            VARCHAR(20)   NOT NULL DEFAULT 'teacher',
    est_actif       BOOLEAN       NOT NULL DEFAULT TRUE,
    created_at      TIMESTAMPTZ   NOT NULL DEFAULT now(),

    CONSTRAINT pk_users       PRIMARY KEY (id),
    CONSTRAINT uq_users_email UNIQUE (email),
    CONSTRAINT ck_users_role  CHECK (role IN ('admin', 'secretary', 'teacher'))
);

CREATE INDEX ix_users_email ON users (email);

COMMENT ON TABLE  users                 IS 'Comptes d''authentification.';
COMMENT ON COLUMN users.hashed_password IS 'Haché bcrypt — jamais en clair.';
COMMENT ON COLUMN users.role            IS 'RBAC : admin | secretary | teacher.';''',
    cols=[
        ('id',              'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY',   'Identifiant auto-incrémenté'),
        ('email',           'VARCHAR(255)', 'NOT NULL, UNIQUE',                   'Adresse courriel (identifiant de connexion)'),
        ('hashed_password', 'VARCHAR(255)', 'NOT NULL',                           'Mot de passe haché bcrypt'),
        ('role',            'VARCHAR(20)',  'NOT NULL, DEFAULT \'teacher\', CHECK', 'Rôle RBAC : admin | secretary | teacher'),
        ('est_actif',       'BOOLEAN',      'NOT NULL, DEFAULT TRUE',             'Compte actif ou désactivé sans suppression'),
        ('created_at',      'TIMESTAMPTZ',  'NOT NULL, DEFAULT now()',            'Horodatage de création du compte'),
    ],
    caption='Tableau 2. Colonnes de la table USERS.',
    index_note='Index sur email — recherche O(log n) à l\'authentification.',
  ),

  # ── SECTION II : TEACHERS ────────────────────────────────────────────────
  dict(
    num='II', title='TABLE TEACHERS — Profils des enseignants',
    description=(
        'La table TEACHERS stocke le profil métier de chaque enseignant. '
        'Elle est liée à la table USERS par une relation 1:1 optionnelle (user_id nullable) : '
        'un enseignant peut être enregistré dans le système sans disposer encore d\'un compte d\'authentification. '
        'Le statut (Permanent ou Vacataire) conditionne le quota horaire annuel applicable, '
        'défini dans la table QUOTAS_STATUTAIRES.'
    ),
    sql='''\
-- ===================================================================
-- 2. ENSEIGNANTS (profil métier)
-- ===================================================================
CREATE TABLE teachers (
    id           INTEGER       GENERATED ALWAYS AS IDENTITY,
    nom          VARCHAR(100)  NOT NULL,
    prenom       VARCHAR(100)  NOT NULL,
    grade        VARCHAR(100)  NOT NULL,
    statut       VARCHAR(20)   NOT NULL,
    departement  VARCHAR(150)  NOT NULL,
    taux_horaire NUMERIC(8,2)  NOT NULL DEFAULT 0.00,
    email        VARCHAR(255)  NOT NULL,
    telephone    VARCHAR(30),
    user_id      INTEGER,

    CONSTRAINT pk_teachers              PRIMARY KEY (id),
    CONSTRAINT uq_teachers_email        UNIQUE (email),
    CONSTRAINT uq_teachers_user         UNIQUE (user_id),
    CONSTRAINT ck_teachers_statut       CHECK (statut IN ('Permanent', 'Vacataire')),
    CONSTRAINT ck_teachers_taux_horaire CHECK (taux_horaire >= 0),
    CONSTRAINT fk_teachers_user         FOREIGN KEY (user_id)
        REFERENCES users (id) ON UPDATE CASCADE ON DELETE SET NULL
);

CREATE INDEX ix_teachers_user_id     ON teachers (user_id);
CREATE INDEX ix_teachers_departement ON teachers (departement);''',
    cols=[
        ('id',           'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY', 'Identifiant auto-incrémenté'),
        ('nom',          'VARCHAR(100)', 'NOT NULL',                          'Nom de famille'),
        ('prenom',       'VARCHAR(100)', 'NOT NULL',                          'Prénom'),
        ('grade',        'VARCHAR(100)', 'NOT NULL',                          'Grade académique (Professeur, Maître-Assistant…)'),
        ('statut',       'VARCHAR(20)',  'NOT NULL, CHECK',                   'Permanent ou Vacataire'),
        ('departement',  'VARCHAR(150)', 'NOT NULL',                          'Département d\'appartenance'),
        ('taux_horaire', 'NUMERIC(8,2)', 'NOT NULL, DEFAULT 0, CHECK ≥ 0',   'Rémunération en FCFA par heure'),
        ('email',        'VARCHAR(255)', 'NOT NULL, UNIQUE',                  'Adresse courriel professionnelle'),
        ('telephone',    'VARCHAR(30)',  'NULL autorisé',                     'Numéro de téléphone (optionnel)'),
        ('user_id',      'INTEGER',      'FK → users.id, UNIQUE, NULL OK',   'Lien vers le compte utilisateur (1:1)'),
    ],
    caption='Tableau 3. Colonnes de la table TEACHERS.',
    index_note='Index sur user_id (jointure fréquente) et departement (filtres statistiques).',
  ),

  # ── SECTION III : COURSES ─────────────────────────────────────────────────
  dict(
    num='III', title='TABLE COURSES — Catalogue des cours',
    description=(
        'La table COURSES constitue le catalogue des unités d\'enseignement (UE) '
        'pour lesquelles des ressources pédagogiques numériques sont produites. '
        'Elle sert de référentiel partagé entre la secrétaire (qui saisit les activités) '
        'et l\'administration. Le niveau est contraint aux cinq paliers du LMD ivoirien.'
    ),
    sql='''\
-- ===================================================================
-- 3. COURS (catalogue pédagogique)
-- ===================================================================
CREATE TABLE courses (
    id         INTEGER       GENERATED ALWAYS AS IDENTITY,
    intitule   VARCHAR(255)  NOT NULL,
    filiere    VARCHAR(150)  NOT NULL,
    niveau     VARCHAR(10)   NOT NULL,
    semestre   VARCHAR(20)   NOT NULL,
    nb_heures  INTEGER,
    nb_credits INTEGER,

    CONSTRAINT pk_courses           PRIMARY KEY (id),
    CONSTRAINT ck_courses_niveau    CHECK (niveau IN ('L1','L2','L3','M1','M2')),
    CONSTRAINT ck_courses_nb_heures CHECK (nb_heures  IS NULL OR nb_heures  >= 0),
    CONSTRAINT ck_courses_credits   CHECK (nb_credits IS NULL OR nb_credits >= 0)
);

CREATE INDEX ix_courses_filiere ON courses (filiere);
CREATE INDEX ix_courses_niveau  ON courses (niveau);''',
    cols=[
        ('id',         'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY', 'Identifiant auto-incrémenté'),
        ('intitule',   'VARCHAR(255)', 'NOT NULL',                          'Intitulé complet du cours'),
        ('filiere',    'VARCHAR(150)', 'NOT NULL',                          'Filière ou spécialité'),
        ('niveau',     'VARCHAR(10)',  'NOT NULL, CHECK',                   'Niveau LMD : L1, L2, L3, M1 ou M2'),
        ('semestre',   'VARCHAR(20)',  'NOT NULL',                          'Semestre (S1, S2, S3…)'),
        ('nb_heures',  'INTEGER',      'NULL OK, CHECK ≥ 0',               'Volume horaire présentiel (optionnel)'),
        ('nb_credits', 'INTEGER',      'NULL OK, CHECK ≥ 0',               'Nombre de crédits ECTS (optionnel)'),
    ],
    caption='Tableau 4. Colonnes de la table COURSES.',
    index_note='Index sur filière et niveau — filtres fréquents dans les listes déroulantes.',
  ),

  # ── SECTION IV : ACADEMIC_YEARS ───────────────────────────────────────────
  dict(
    num='IV', title='TABLE ACADEMIC_YEARS — Années académiques',
    description=(
        'La table ACADEMIC_YEARS permet de rattacher chaque activité à une année académique précise, '
        'facilitant les calculs de volumes horaires annuels et les comparaisons inter-années. '
        'Un index partiel unique garantit qu\'une seule année peut être marquée comme active (status = TRUE) '
        'à un instant donné, sans nécessiter de trigger ni de logique applicative complexe.'
    ),
    sql='''\
-- ===================================================================
-- 4. ANNÉES ACADÉMIQUES
-- ===================================================================
CREATE TABLE academic_years (
    id         INTEGER      GENERATED ALWAYS AS IDENTITY,
    libelle    VARCHAR(20)  NOT NULL,
    date_debut DATE,
    date_fin   DATE,
    status     BOOLEAN      NOT NULL DEFAULT FALSE,

    CONSTRAINT pk_academic_years       PRIMARY KEY (id),
    CONSTRAINT uq_academic_years_lib   UNIQUE (libelle),
    CONSTRAINT ck_academic_years_dates CHECK (
        date_debut IS NULL OR date_fin IS NULL OR date_fin >= date_debut
    )
);

-- Une seule année active à la fois (index partiel unique)
CREATE UNIQUE INDEX uq_academic_years_active
    ON academic_years (status) WHERE status = TRUE;''',
    cols=[
        ('id',         'INTEGER',     'PK, GENERATED ALWAYS AS IDENTITY', 'Identifiant auto-incrémenté'),
        ('libelle',    'VARCHAR(20)', 'NOT NULL, UNIQUE',                  'Libellé, ex. « 2025-2026 »'),
        ('date_debut', 'DATE',        'NULL OK',                           'Date de début de l\'année'),
        ('date_fin',   'DATE',        'NULL OK, CHECK ≥ date_debut',      'Date de fin de l\'année'),
        ('status',     'BOOLEAN',     'NOT NULL, DEFAULT FALSE',           'TRUE = année courante active (une seule à la fois)'),
    ],
    caption='Tableau 5. Colonnes de la table ACADEMIC_YEARS.',
    index_note='Index partiel unique sur status WHERE status = TRUE — unicité de l\'année active sans colonne redondante.',
  ),

  # ── SECTION V : RESOURCES ─────────────────────────────────────────────────
  dict(
    num='V', title='TABLE RESOURCES — Ressources pédagogiques',
    description=(
        'La table RESOURCES représente la combinaison unique '
        '(enseignant × cours × type d\'activité × niveau de complexité). '
        'Elle constitue la couche de référentiel entre les entités métier et les activités déclarées. '
        'La contrainte d\'unicité composite sur (teacher_id, course_id, type, niveau_complexite) '
        'empêche la création de doublons de ressources pour un même enseignant.'
    ),
    sql='''\
-- ===================================================================
-- 5. RESSOURCES PÉDAGOGIQUES
-- ===================================================================
CREATE TABLE resources (
    id                INTEGER      GENERATED ALWAYS AS IDENTITY,
    type              VARCHAR(30)  NOT NULL,
    niveau_complexite SMALLINT     NOT NULL DEFAULT 1,
    course_id         INTEGER      NOT NULL,
    teacher_id        INTEGER      NOT NULL,
    date_creation     TIMESTAMPTZ  NOT NULL DEFAULT now(),

    CONSTRAINT pk_resources          PRIMARY KEY (id),
    CONSTRAINT uq_resources_business UNIQUE (teacher_id, course_id, type, niveau_complexite),
    CONSTRAINT ck_resources_type     CHECK (type IN ('creation', 'mise_a_jour')),
    CONSTRAINT ck_resources_niveau   CHECK (niveau_complexite IN (1, 2, 3)),
    CONSTRAINT fk_resources_course   FOREIGN KEY (course_id)
        REFERENCES courses (id)  ON UPDATE CASCADE ON DELETE RESTRICT,
    CONSTRAINT fk_resources_teacher  FOREIGN KEY (teacher_id)
        REFERENCES teachers (id) ON UPDATE CASCADE ON DELETE RESTRICT
);

CREATE INDEX ix_resources_course_id  ON resources (course_id);
CREATE INDEX ix_resources_teacher_id ON resources (teacher_id);''',
    cols=[
        ('id',                'INTEGER',     'PK, GENERATED ALWAYS AS IDENTITY', 'Identifiant auto-incrémenté'),
        ('type',              'VARCHAR(30)', 'NOT NULL, CHECK',                   'creation ou mise_a_jour'),
        ('niveau_complexite', 'SMALLINT',    'NOT NULL, DEFAULT 1, CHECK 1–3',   'Niveau 1, 2 ou 3 (détermine le coefficient)'),
        ('course_id',         'INTEGER',     'NOT NULL, FK → courses.id',        'Cours concerné par la ressource'),
        ('teacher_id',        'INTEGER',     'NOT NULL, FK → teachers.id',       'Enseignant auteur de la ressource'),
        ('date_creation',     'TIMESTAMPTZ', 'NOT NULL, DEFAULT now()',           'Horodatage de création'),
    ],
    caption='Tableau 6. Colonnes de la table RESOURCES.',
    index_note='Contrainte UNIQUE composite (teacher_id, course_id, type, niveau) garantit l\'unicité métier.',
  ),

  # ── SECTION VI : ACTIVITIES ───────────────────────────────────────────────
  dict(
    num='VI', title='TABLE ACTIVITIES — Activités pédagogiques déclarées',
    description=(
        'La table ACTIVITIES enregistre chaque acte de production pédagogique '
        'générateur d\'un volume horaire calculé (Vhtc). '
        'La formule officielle UVCI est : Vhtc = Ic × S, '
        'où Ic est le coefficient de complexité (issu de COEFFICIENT_CONFIGS) '
        'et S est le nombre de séquences produites. '
        'Les activités suivent un cycle de validation à trois états : '
        'en_attente → valide | rejetee, avec traçabilité du validateur.'
    ),
    sql='''\
-- ===================================================================
-- 6. ACTIVITÉS (volume horaire généré)
-- ===================================================================
CREATE TABLE activities (
    id                     INTEGER       GENERATED ALWAYS AS IDENTITY,
    type                   VARCHAR(20)   NOT NULL,
    resource_id            INTEGER       NOT NULL,
    teacher_id             INTEGER       NOT NULL,
    nb_sequences           INTEGER       NOT NULL DEFAULT 1,
    volume_horaire_calcule NUMERIC(8,3)  NOT NULL DEFAULT 0.000,
    academic_year_id       INTEGER,
    annee_academique       VARCHAR(20),
    validation_status      VARCHAR(20)   NOT NULL DEFAULT 'en_attente',
    validated_by           INTEGER,
    validated_at           TIMESTAMPTZ,
    created_at             TIMESTAMPTZ   NOT NULL DEFAULT now(),

    CONSTRAINT pk_activities           PRIMARY KEY (id),
    CONSTRAINT ck_activities_type      CHECK (type IN ('creation', 'mise_a_jour')),
    CONSTRAINT ck_activities_status    CHECK (validation_status IN
                                        ('en_attente', 'valide', 'rejetee')),
    CONSTRAINT ck_activities_sequences CHECK (nb_sequences > 0),
    CONSTRAINT ck_activities_volume    CHECK (volume_horaire_calcule >= 0),
    CONSTRAINT fk_activities_resource  FOREIGN KEY (resource_id)
        REFERENCES resources (id)      ON UPDATE CASCADE ON DELETE CASCADE,
    CONSTRAINT fk_activities_teacher   FOREIGN KEY (teacher_id)
        REFERENCES teachers (id)       ON UPDATE CASCADE ON DELETE RESTRICT,
    CONSTRAINT fk_activities_year      FOREIGN KEY (academic_year_id)
        REFERENCES academic_years (id) ON UPDATE CASCADE ON DELETE SET NULL,
    CONSTRAINT fk_activities_validator FOREIGN KEY (validated_by)
        REFERENCES users (id)          ON UPDATE CASCADE ON DELETE SET NULL
);

CREATE INDEX ix_activities_resource_id ON activities (resource_id);
CREATE INDEX ix_activities_teacher_id  ON activities (teacher_id);
CREATE INDEX ix_activities_year_id     ON activities (academic_year_id);
CREATE INDEX ix_activities_status      ON activities (validation_status);''',
    cols=[
        ('id',                     'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY',     'Identifiant auto-incrémenté'),
        ('type',                   'VARCHAR(20)',  'NOT NULL, CHECK',                        'creation (Ic plein) ou mise_a_jour (½ Ic)'),
        ('resource_id',            'INTEGER',      'NOT NULL, FK → resources.id, CASCADE', 'Ressource concernée (suppression en cascade)'),
        ('teacher_id',             'INTEGER',      'NOT NULL, FK → teachers.id, RESTRICT', 'Enseignant déclarant l\'activité'),
        ('nb_sequences',           'INTEGER',      'NOT NULL, DEFAULT 1, CHECK > 0',        'Nombre de séquences S produites'),
        ('volume_horaire_calcule', 'NUMERIC(8,3)', 'NOT NULL, DEFAULT 0, CHECK ≥ 0',       'Vhtc = Ic × S (3 décimales)'),
        ('academic_year_id',       'INTEGER',      'NULL OK, FK → academic_years.id',      'Année académique (NULL si non renseignée)'),
        ('annee_academique',       'VARCHAR(20)',  'NULL OK',                               'Libellé textuel complémentaire (dénormalisation)'),
        ('validation_status',      'VARCHAR(20)',  'NOT NULL, DEFAULT en_attente, CHECK',  'État : en_attente | valide | rejetee'),
        ('validated_by',           'INTEGER',      'NULL OK, FK → users.id',               'Compte admin/secrétaire ayant validé'),
        ('validated_at',           'TIMESTAMPTZ',  'NULL OK',                               'Horodatage de la validation'),
        ('created_at',             'TIMESTAMPTZ',  'NOT NULL, DEFAULT now()',               'Horodatage de création'),
    ],
    caption='Tableau 7. Colonnes de la table ACTIVITIES.',
    index_note='4 index : resource_id, teacher_id, academic_year_id, validation_status — colonnes de filtrage fréquent.',
  ),

  # ── SECTION VII : COEFFICIENT_CONFIGS ────────────────────────────────────
  dict(
    num='VII', title='TABLE COEFFICIENT_CONFIGS — Barème officiel des coefficients',
    description=(
        'La table COEFFICIENT_CONFIGS stocke le barème officiel UVCI des coefficients horaires (Ic) '
        'utilisés dans la formule Vhtc = Ic × S. '
        'L\'administrateur peut modifier ces valeurs via l\'API (PUT /config/coefficients) '
        'sans redéploiement. Les six lignes initiales correspondent au barème officiel, '
        'avec le principe que le coefficient de mise à jour vaut la moitié de celui de création.'
    ),
    sql='''\
-- ===================================================================
-- 7. CONFIGURATION DES COEFFICIENTS (barème UVCI)
-- ===================================================================
CREATE TABLE coefficient_configs (
    id                INTEGER       GENERATED ALWAYS AS IDENTITY,
    niveau_complexite SMALLINT      NOT NULL,
    type_activite     VARCHAR(20)   NOT NULL,
    coefficient       NUMERIC(6,3)  NOT NULL,

    CONSTRAINT pk_coefficient_configs PRIMARY KEY (id),
    CONSTRAINT uq_coeff_niveau_type   UNIQUE (niveau_complexite, type_activite),
    CONSTRAINT ck_coeff_niveau        CHECK (niveau_complexite IN (1, 2, 3)),
    CONSTRAINT ck_coeff_type          CHECK (type_activite IN ('creation', 'mise_a_jour')),
    CONSTRAINT ck_coeff_value         CHECK (coefficient >= 0)
);''',
    cols=[
        ('id',                'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY',       'Identifiant auto-incrémenté'),
        ('niveau_complexite', 'SMALLINT',     'NOT NULL, UNIQUE composite, CHECK 1–3',  'Niveau de complexité du contenu (1, 2 ou 3)'),
        ('type_activite',     'VARCHAR(20)',  'NOT NULL, UNIQUE composite, CHECK',      'creation ou mise_a_jour'),
        ('coefficient',       'NUMERIC(6,3)', 'NOT NULL, CHECK ≥ 0',                   'Valeur du coefficient Ic (3 décimales)'),
    ],
    caption='Tableau 8. Colonnes de la table COEFFICIENT_CONFIGS.',
    index_note='Contrainte UNIQUE (niveau_complexite, type_activite) — une seule ligne par combinaison niveau × type.',
  ),

  # ── SECTION VIII : QUOTAS_STATUTAIRES ────────────────────────────────────
  dict(
    num='VIII', title='TABLE QUOTAS_STATUTAIRES — Quotas horaires annuels',
    description=(
        'La table QUOTAS_STATUTAIRES définit le volume horaire annuel dû '
        'pour chaque combinaison (grade × statut). '
        'Ces valeurs permettent de calculer, pour chaque enseignant, '
        'la proportion du quota atteint (tableau de bord). '
        'L\'administrateur peut les ajuster sans redéploiement via l\'API (PUT /config/quotas). '
        'Par défaut : 192 h pour les permanents, 96 h pour les vacataires, quel que soit le grade.'
    ),
    sql='''\
-- ===================================================================
-- 8. QUOTAS STATUTAIRES
-- ===================================================================
CREATE TABLE quotas_statutaires (
    id           INTEGER       GENERATED ALWAYS AS IDENTITY,
    grade        VARCHAR(100)  NOT NULL,
    statut       VARCHAR(20)   NOT NULL,
    quota_heures NUMERIC(8,2)  NOT NULL,

    CONSTRAINT pk_quotas_statutaires PRIMARY KEY (id),
    CONSTRAINT uq_quota_grade_statut UNIQUE (grade, statut),
    CONSTRAINT ck_quota_statut       CHECK (statut IN ('Permanent', 'Vacataire')),
    CONSTRAINT ck_quota_value        CHECK (quota_heures >= 0)
);''',
    cols=[
        ('id',           'INTEGER',      'PK, GENERATED ALWAYS AS IDENTITY',  'Identifiant auto-incrémenté'),
        ('grade',        'VARCHAR(100)', 'NOT NULL, UNIQUE composite',        'Grade académique de l\'enseignant'),
        ('statut',       'VARCHAR(20)',  'NOT NULL, UNIQUE composite, CHECK', 'Permanent ou Vacataire'),
        ('quota_heures', 'NUMERIC(8,2)', 'NOT NULL, CHECK ≥ 0',              'Quota horaire annuel en heures'),
    ],
    caption='Tableau 9. Colonnes de la table QUOTAS_STATUTAIRES.',
    index_note='Contrainte UNIQUE (grade, statut) — une seule ligne par combinaison grade × statut.',
  ),
]

# ── rendu des sections ──────────────────────────────────────────────────────
tbl_counter = 2  # déjà Tableau 1 et 2
for s in sections:
    section_title(doc, s['num'], s['title'])
    hr(doc)
    add_text(doc, s['description'])
    subsection_title(doc, f'Script DDL — SECTION {s["num"]}')
    add_sql_block(doc, s['sql'])
    subsection_title(doc, 'Description des colonnes')
    add_col_table(doc, s['title'], s['cols'], caption=s['caption'])
    p_idx = doc.add_paragraph()
    p_idx.paragraph_format.space_before = Pt(4)
    p_idx.paragraph_format.space_after  = Pt(8)
    ri = p_idx.add_run(f'Note : {s["index_note"]}')
    ri.font.italic = True
    ri.font.size   = Pt(10)
    ri.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


# ════════════════════════════════════════════════════════════════════════════
# SECTION IX : DONNÉES DE RÉFÉRENCE
# ════════════════════════════════════════════════════════════════════════════

section_title(doc, 'IX', 'DONNÉES DE RÉFÉRENCE — Instructions INSERT')
hr(doc)
add_text(doc, (
    'Cette section regroupe les données initiales insérées au moment de la création du schéma. '
    'Elles correspondent au barème officiel UVCI des coefficients horaires (six lignes) '
    'et aux quotas statutaires par défaut (huit lignes, pour quatre grades × deux statuts). '
    'Ces valeurs peuvent être modifiées ultérieurement par l\'administrateur via l\'interface.'
))

subsection_title(doc, 'IX.1  Barème officiel des coefficients horaires')
add_sql_block(doc, '''\
-- Barème officiel UVCI : Vhtc = Ic × S ; mise à jour = ½ × création.
INSERT INTO coefficient_configs (niveau_complexite, type_activite, coefficient) VALUES
    (1, 'creation',    0.400),   -- Niveau 1 : contenus simples + quiz
    (1, 'mise_a_jour', 0.200),
    (2, 'creation',    0.750),   -- Niveau 2 : +25 % activités interactives
    (2, 'mise_a_jour', 0.375),
    (3, 'creation',    1.500),   -- Niveau 3 : serious games, simulations
    (3, 'mise_a_jour', 0.750);''')

# Tableau barème
bareme_rows = [
    ('Niveau 1', 'creation',    '0.400', 'Contenus simples + quiz'),
    ('Niveau 1', 'mise_a_jour', '0.200', '½ × création (mise à jour)'),
    ('Niveau 2', 'creation',    '0.750', '+25 % activités interactives'),
    ('Niveau 2', 'mise_a_jour', '0.375', '½ × création (mise à jour)'),
    ('Niveau 3', 'creation',    '1.500', 'Serious games, simulations'),
    ('Niveau 3', 'mise_a_jour', '0.750', '½ × création (mise à jour)'),
]
bt = doc.add_table(rows=1 + len(bareme_rows), cols=4)
bt.style = 'Table Grid'
for i, h in enumerate(['Niveau', 'Type', 'Coefficient Ic', 'Description']):
    cell = bt.cell(0, i)
    shade_cell(cell, GREY10)
    r = cell.paragraphs[0].add_run(h)
    set_font(r, 10, bold=True, color=NAVY)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

widths_b = [3.0, 3.5, 3.5, 6.5]
for i, w in enumerate(widths_b):
    set_col_width(bt, i, w)

for ri, (niv, typ, coef, desc) in enumerate(bareme_rows):
    bg = GREY05 if ri % 2 else WHITE
    for ci, val in enumerate([niv, typ, coef, desc]):
        cell = bt.cell(ri + 1, ci)
        if bg != WHITE:
            shade_cell(cell, bg)
        r = cell.paragraphs[0].add_run(val)
        fname = 'Courier New' if ci in (1, 2) else 'Times New Roman'
        bold = (ci == 2)
        set_font(r, 9.5, bold=bold, name=fname,
                 color=NAVY if ci == 2 else BLACK)
        if ci == 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

cp = doc.add_paragraph('Tableau 10. Barème officiel UVCI des coefficients horaires.')
cp.paragraph_format.space_before = Pt(2)
cp.paragraph_format.space_after  = Pt(10)
cp.runs[0].font.italic = True
cp.runs[0].font.size   = Pt(10)
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

subsection_title(doc, 'IX.2  Quotas statutaires par défaut')
add_sql_block(doc, '''\
-- Quotas statutaires par défaut : 192 h (permanents), 96 h (vacataires).
INSERT INTO quotas_statutaires (grade, statut, quota_heures) VALUES
    ('Professeur',            'Permanent', 192.00),
    ('Maître de Conférences', 'Permanent', 192.00),
    ('Maître-Assistant',      'Permanent', 192.00),
    ('Assistant',             'Permanent', 192.00),
    ('Professeur',            'Vacataire',  96.00),
    ('Maître de Conférences', 'Vacataire',  96.00),
    ('Maître-Assistant',      'Vacataire',  96.00),
    ('Assistant',             'Vacataire',  96.00);''')

quota_rows = [
    ('Professeur',            'Permanent', '192.00'),
    ('Maître de Conférences', 'Permanent', '192.00'),
    ('Maître-Assistant',      'Permanent', '192.00'),
    ('Assistant',             'Permanent', '192.00'),
    ('Professeur',            'Vacataire',  '96.00'),
    ('Maître de Conférences', 'Vacataire',  '96.00'),
    ('Maître-Assistant',      'Vacataire',  '96.00'),
    ('Assistant',             'Vacataire',  '96.00'),
]
qt = doc.add_table(rows=1 + len(quota_rows), cols=3)
qt.style = 'Table Grid'
for i, h in enumerate(['Grade', 'Statut', 'Quota (heures/an)']):
    cell = qt.cell(0, i)
    shade_cell(cell, GREY10)
    r = cell.paragraphs[0].add_run(h)
    set_font(r, 10, bold=True, color=NAVY)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, w in enumerate([6.5, 4.0, 4.0]):
    set_col_width(qt, i, w)

for ri, (grade, statut, quota) in enumerate(quota_rows):
    bg = GREY05 if ri % 2 else WHITE
    for ci, val in enumerate([grade, statut, quota]):
        cell = qt.cell(ri + 1, ci)
        if bg != WHITE:
            shade_cell(cell, bg)
        r = cell.paragraphs[0].add_run(val)
        set_font(r, 9.5, bold=(ci == 2),
                 color=NAVY if ci == 2 else BLACK)
        if ci == 2:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

cp2 = doc.add_paragraph('Tableau 11. Quotas statutaires par défaut (modifiables via l\'interface admin).')
cp2.paragraph_format.space_before = Pt(2)
cp2.paragraph_format.space_after  = Pt(10)
cp2.runs[0].font.italic = True
cp2.runs[0].font.size   = Pt(10)
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
# ANNEXE A : RÉCAPITULATIF DES TABLES
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.page_break_before = True
r = p.add_run('ANNEXE A. RÉCAPITULATIF DES TABLES')
set_font(r, 13, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
hr(doc)

recap_rows = [
    ('users',               '6',  'Comptes d\'authentification (admin / secretary / teacher)'),
    ('teachers',            '10', 'Profils métier des enseignants'),
    ('courses',             '7',  'Catalogue des unités d\'enseignement'),
    ('academic_years',      '5',  'Années académiques de rattachement'),
    ('resources',           '6',  'Combinaisons enseignant × cours × type × niveau'),
    ('activities',          '12', 'Activités déclarées et volumes horaires calculés'),
    ('coefficient_configs', '4',  'Barème officiel des coefficients Ic (paramétrable)'),
    ('quotas_statutaires',  '4',  'Quotas horaires annuels par grade et statut (paramétrable)'),
]
rt = doc.add_table(rows=1 + len(recap_rows), cols=3)
rt.style = 'Table Grid'
for i, h in enumerate(['Table', 'Colonnes', 'Rôle']):
    cell = rt.cell(0, i)
    shade_cell(cell, GREY10)
    r = cell.paragraphs[0].add_run(h)
    set_font(r, 10, bold=True, color=NAVY)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, w in enumerate([4.5, 2.5, 9.5]):
    set_col_width(rt, i, w)

for ri, (tname, ncols, role) in enumerate(recap_rows):
    bg = GREY05 if ri % 2 else WHITE
    for ci, val in enumerate([tname, ncols, role]):
        cell = rt.cell(ri + 1, ci)
        if bg != WHITE:
            shade_cell(cell, bg)
        r = cell.paragraphs[0].add_run(val)
        set_font(r, 9.5, name='Courier New' if ci == 0 else 'Times New Roman',
                 bold=(ci == 0), color=NAVY if ci == 0 else BLACK)
        if ci == 1:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

cpa = doc.add_paragraph('Tableau 12. Récapitulatif des 8 tables du schéma UVCI.')
cpa.paragraph_format.space_before = Pt(2)
cpa.paragraph_format.space_after  = Pt(10)
cpa.runs[0].font.italic = True
cpa.runs[0].font.size   = Pt(10)
cpa.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
# ANNEXE B : MATRICE DES CONTRAINTES
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
r = p.add_run('ANNEXE B. MATRICE DES CONTRAINTES D\'INTÉGRITÉ')
set_font(r, 13, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
hr(doc)

mat_rows = [
    ('users',               '1 PK',  '0 FK', '1 UNIQUE', '1 CHECK'),
    ('teachers',            '1 PK',  '1 FK', '2 UNIQUE', '2 CHECK'),
    ('courses',             '1 PK',  '0 FK', '0 UNIQUE', '3 CHECK'),
    ('academic_years',      '1 PK',  '0 FK', '2 UNIQUE', '1 CHECK'),
    ('resources',           '1 PK',  '2 FK', '1 UNIQUE', '2 CHECK'),
    ('activities',          '1 PK',  '4 FK', '0 UNIQUE', '4 CHECK'),
    ('coefficient_configs', '1 PK',  '0 FK', '1 UNIQUE', '3 CHECK'),
    ('quotas_statutaires',  '1 PK',  '0 FK', '1 UNIQUE', '2 CHECK'),
    ('TOTAL',               '8',     '7',    '8',         '18'),
]
mt = doc.add_table(rows=1 + len(mat_rows), cols=5)
mt.style = 'Table Grid'
for i, h in enumerate(['Table', 'Clés primaires', 'Clés étrangères', 'Contraintes UNIQUE', 'Contraintes CHECK']):
    cell = mt.cell(0, i)
    shade_cell(cell, GREY10)
    r = cell.paragraphs[0].add_run(h)
    set_font(r, 10, bold=True, color=NAVY)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, w in enumerate([4.5, 3.0, 3.0, 3.5, 3.0]):
    set_col_width(mt, i, w)

for ri, row in enumerate(mat_rows):
    is_total = (row[0] == 'TOTAL')
    bg = GREY10 if is_total else (GREY05 if ri % 2 else WHITE)
    for ci, val in enumerate(row):
        cell = mt.cell(ri + 1, ci)
        shade_cell(cell, bg)
        r = cell.paragraphs[0].add_run(val)
        set_font(r, 9.5, bold=is_total,
                 name='Courier New' if ci == 0 else 'Times New Roman',
                 color=NAVY if (ci == 0 or is_total) else BLACK)
        if ci > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

cpb = doc.add_paragraph('Tableau 13. Matrice récapitulative des contraintes d\'intégrité par table.')
cpb.paragraph_format.space_before = Pt(2)
cpb.paragraph_format.space_after  = Pt(10)
cpb.runs[0].font.italic = True
cpb.runs[0].font.size   = Pt(10)
cpb.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ════════════════════════════════════════════════════════════════════════════
# ANNEXE C : SCHÉMA DES RELATIONS
# ════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(16)
r = p.add_run('ANNEXE C. SCHÉMA DES RELATIONS (SYNTHÈSE)')
set_font(r, 13, bold=True, color=NAVY)
p.paragraph_format.space_after = Pt(6)
hr(doc)

add_text(doc, (
    'Le diagramme textuel ci-dessous synthétise les relations entre les huit tables du schéma. '
    'Les flèches indiquent le sens de la clé étrangère (table enfant → table parent). '
    'Les cardinalités sont précisées pour chaque relation.'
))

rel_sql = '''\
users (1) ────────────────── (0,1) teachers
    │                                │
    │ validated_by                   │ (1,n)
    ▼                                ▼
activities (n) ──── (1) resources (n) ──── (1) courses
    │                    │
    │ academic_year_id   │ teacher_id (redondant, optimisation)
    ▼                    │
academic_years (1)       │
                         ▼
                    teachers (1)

-- Tables de paramétrage (sans FK vers les tables métier) :
coefficient_configs   →  utilisée par le service de calcul (Vhtc = Ic × S)
quotas_statutaires    →  utilisée par le tableau de bord (taux d'atteinte du quota)'''

add_sql_block(doc, rel_sql)

add_text(doc, (
    'La table RESOURCES joue le rôle de table de jonction entre TEACHERS, COURSES '
    'et le type d\'activité/niveau, matérialisant la contrainte métier d\'unicité. '
    'La table ACTIVITIES référence RESOURCES (et non directement TEACHERS × COURSES) '
    'afin de garantir la cohérence des données et d\'éviter les duplications.'
), sb=4)

# Espace final
doc.add_paragraph()

# ─── sauvegarde ──────────────────────────────────────────────────────────────
out = '/Users/apple/projet/projet_UVCI/Script_BD_UVCI.docx'
doc.save(out)
print(f'✓  Sauvegardé → {out}')
